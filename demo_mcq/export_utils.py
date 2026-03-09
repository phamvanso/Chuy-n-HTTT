"""
export_utils.py
───────────────
Xuất đề thi MCQ ra Word (.docx) và PDF.

Dùng:
    from export_utils import export_word_bytes, export_pdf_bytes
    docx_bytes = export_word_bytes(mcq_list, title="Đề kiểm tra")
    pdf_bytes  = export_pdf_bytes(mcq_list, title="Đề kiểm tra")
"""

from __future__ import annotations
from typing import List, Dict

LABELS = ["A", "B", "C", "D", "E"]


# ══════════════════════════════════════════════════════════════
# WORD EXPORT  (python-docx)
# ══════════════════════════════════════════════════════════════
def export_word_bytes(mcq_list: List[Dict], title: str = "Đề kiểm tra trắc nghiệm") -> bytes:
    """Trả về bytes của file .docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("Chạy: pip install python-docx")

    doc = Document()

    # ── Tiêu đề ────────────────────────────────────────────────
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # khoảng trắng

    # ── Nội dung câu hỏi ───────────────────────────────────────
    for i, mcq in enumerate(mcq_list, 1):
        # Câu hỏi
        q_para = doc.add_paragraph()
        q_run  = q_para.add_run(f"Câu {i}. {mcq['question']}")
        q_run.bold = True
        q_run.font.size = Pt(12)

        # Các lựa chọn
        for j, opt in enumerate(mcq["options"]):
            lbl = LABELS[j]
            is_correct = (lbl == mcq["correct_label"])
            opt_para = doc.add_paragraph(style="List Bullet")
            opt_para.paragraph_format.left_indent = Inches(0.3)
            run = opt_para.add_run(f"{lbl}. {opt}")
            run.font.size = Pt(11)
            if is_correct:
                run.bold = True
                run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)  # xanh lá

        doc.add_paragraph("")  # khoảng cách giữa các câu

    # ── Đáp án (trang riêng) ────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("ĐÁP ÁN", level=2)
    ans_lines = []
    for i, mcq in enumerate(mcq_list, 1):
        ans_lines.append(f"Câu {i}: {mcq['correct_label']}")

    # 3 cột đáp án
    cols = [ans_lines[k::3] for k in range(3)]
    max_rows = max(len(c) for c in cols)
    table = doc.add_table(rows=max_rows, cols=3)
    table.style = "Table Grid"
    for col_idx, col_data in enumerate(cols):
        for row_idx, cell_text in enumerate(col_data):
            table.cell(row_idx, col_idx).text = cell_text

    # ── Serialize ──────────────────────────────────────────────
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# PDF EXPORT  (fpdf2)
# ══════════════════════════════════════════════════════════════
def _register_unicode_font(pdf) -> str:
    """
    Đăng ký font TTF hỗ trợ tiếng Việt cho fpdf2.
    Trả về tên family đã đăng ký (dùng cho set_font).
    """
    import os
    candidates = [
        # Windows – Arial
        ("Arial",   r"C:\Windows\Fonts\arial.ttf",   r"C:\Windows\Fonts\arialbd.ttf"),
        # Windows – Calibri
        ("Calibri", r"C:\Windows\Fonts\calibri.ttf", r"C:\Windows\Fonts\calibrib.ttf"),
        # Windows – Times New Roman
        ("Times",   r"C:\Windows\Fonts\times.ttf",   r"C:\Windows\Fonts\timesbd.ttf"),
        # Linux/macOS – DejaVu
        ("DejaVu",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        # macOS
        ("Arial",   "/Library/Fonts/Arial.ttf",      "/Library/Fonts/Arial Bold.ttf"),
    ]
    for family, reg, bold in candidates:
        if os.path.exists(reg):
            pdf.add_font(family, style="", fname=reg)
            pdf.add_font(family, style="B", fname=bold if os.path.exists(bold) else reg)
            return family
    # Không tìm thấy TTF – dùng Helvetica (non-Unicode, dấu tiếng Việt sẽ bị lỗi)
    return "Helvetica"


def export_pdf_bytes(mcq_list: List[Dict], title: str = "Đề kiểm tra trắc nghiệm") -> bytes:
    """Trả về bytes của file .pdf (UTF-8 tiếng Việt)."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("Chạy: pip install fpdf2")

    _title_ref = [title]  # closure-friendly

    class PDF(FPDF):
        def header(self):
            self.set_font(_font, "B", 14)
            self.cell(0, 10, _title_ref[0], align="C", new_x="LMARGIN", new_y="NEXT")
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font(_font, "", 8)
            self.cell(0, 10, f"Trang {self.page_no()}", align="C")

    pdf = PDF()
    _font = _register_unicode_font(pdf)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    for i, mcq in enumerate(mcq_list, 1):
        # Câu hỏi
        pdf.set_font(_font, "B", 12)
        pdf.multi_cell(0, 8, f"Cau {i}. {mcq['question']}", new_x="LMARGIN", new_y="NEXT")

        # Lựa chọn
        for j, opt in enumerate(mcq["options"]):
            lbl = LABELS[j]
            is_correct = (lbl == mcq["correct_label"])
            if is_correct:
                pdf.set_text_color(22, 163, 74)
                pdf.set_font(_font, "B", 11)
            else:
                pdf.set_text_color(55, 65, 81)
                pdf.set_font(_font, "", 11)
            pdf.multi_cell(0, 7, f"   {lbl}. {opt}", new_x="LMARGIN", new_y="NEXT")

        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

    # Trang đáp án
    pdf.add_page()
    pdf.set_font(_font, "B", 13)
    pdf.cell(0, 10, "DAP AN", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font(_font, "", 11)
    for i, mcq in enumerate(mcq_list, 1):
        pdf.cell(60, 8, f"Cau {i}: {mcq['correct_label']}")
        if i % 3 == 0:
            pdf.ln(8)
    pdf.ln(8)

    import io
    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()
